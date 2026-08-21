from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

out = Path(__file__).with_name('automation-testing-guide.docx')
d = Document(); s = d.sections[0]
s.top_margin=Inches(.7); s.bottom_margin=Inches(.7); s.left_margin=Inches(.8); s.right_margin=Inches(.8)
d.styles['Normal'].font.name='Aptos'; d.styles['Normal'].font.size=Pt(10)
for n,z,c in [('Title',24,'123B5D'),('Heading 1',16,'123B5D'),('Heading 2',12,'1B6B73')]:
    st=d.styles[n]; st.font.name='Aptos Display'; st.font.size=Pt(z); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(c)
def tbl(head, rows):
    t=d.add_table(rows=1, cols=len(head)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Light Shading Accent 1'
    for i,x in enumerate(head):
        c=t.rows[0].cells[i]; c.text=x; sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'123B5D'); c._tc.get_or_add_tcPr().append(sh); c.paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255); c.paragraphs[0].runs[0].font.bold=True
    for row in rows:
        for i,x in enumerate(row): t.add_row().cells[i].text=str(x)
    return t
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('HOTEL SAAS PLATFORM').bold=True
p=d.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Automation Testing & Quality Gates')
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Chiến lược kiểm thử hiện đại, vận hành CI và hướng dẫn phát triển').italic=True
d.add_paragraph('Phiên bản 1.0 | 20/08/2026', style='Subtitle')
d.add_heading('1. Mục tiêu và phạm vi',1); d.add_paragraph('Tài liệu mô tả cách tích hợp và vận hành automation test cho backend .NET 10, frontend Angular, API contract, Playwright E2E, accessibility/performance audit và quality gates trên GitHub Actions.')
d.add_heading('2. Kiến trúc kiểm thử',1); tbl(['Tầng','Công cụ','Phạm vi','Tần suất'],[['Unit','Vitest, xUnit','Hàm, service, policy','Commit/PR'],['Integration','xUnit + EF InMemory','API, tenant isolation, RBAC','PR'],['Contract','Vitest HTTP','DTO, URL, action mask','PR'],['E2E smoke','Playwright Chromium','Search, booking, navigation','PR'],['E2E integration','Playwright','Payment, refund, admin','Nightly/release']])
d.add_heading('3. Lệnh chạy chuẩn',1)
for x in ['npm ci','npm run test:unit:ci','npm run e2e:smoke','npm run test:all','dotnet test backend/HotelSaas.slnx --configuration Release --collect:"XPlat Code Coverage"']: d.add_paragraph(x,style='List Bullet')
d.add_paragraph('Unit chạy không watch và xuất coverage. Smoke E2E dùng cấu hình độc lập; integration E2E cần backend và seed data.')
d.add_heading('4. CI quality gates',1); d.add_paragraph('Workflow .github/workflows/quality-gates.yml chạy ba job song song: backend restore/build/test với TRX và coverage; frontend production build/unit coverage; Playwright cài Chromium và chạy smoke. Artifact được lưu cả khi lỗi để điều tra.')
d.add_heading('5. Quy ước test hiện đại',1)
for x in ['AAA; mỗi test xác nhận một hành vi chính.','Test độc lập, tenant-scoped, không phụ thuộc thứ tự.','Playwright ưu tiên locator role/label/test id.','Bao phủ happy path, validation, authorization, isolation, idempotency và retry.','Giữ trace on-first-retry và screenshot on-failure; không commit artifact/secret.']: d.add_paragraph(x,style='List Bullet')
d.add_heading('6. Ma trận release',1); tbl(['Gate','Dev','PR','Nightly','Release'],[['Backend unit/integration','Có','Có','Có','Có'],['Frontend unit','Có','Có','Có','Có'],['Production build','Tuỳ chọn','Có','Có','Có'],['Playwright smoke','Tuỳ chọn','Có','Có','Có'],['Playwright integration','Không','Không','Có','Có'],['Accessibility/performance','Không','Tuỳ chọn','Có','Có']])
d.add_heading('7. Xử lý lỗi và mở rộng',1)
for x in ['Đọc TRX/coverage hoặc Playwright trace trước khi rerun.','Flaky test phải tái hiện ít nhất ba lần; không tăng retry để che lỗi.','Thêm Testcontainers PostgreSQL/Redis, JUnit/LCOV threshold, OWASP smoke và Lighthouse budget theo lộ trình.']: d.add_paragraph(x,style='List Bullet')
d.add_heading('8. Checklist bàn giao',1)
for x in ['dotnet test pass','npm run test:unit:ci pass','npm run e2e:smoke pass','Production build pass','Không có .only hoặc selector không ổn định','CI lưu artifact khi thất bại']: d.add_paragraph('[ ] '+x)
d.sections[0].footer.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; d.sections[0].footer.paragraphs[0].add_run('Hotel SaaS Platform | Automation Testing Guide')
d.save(out); print(out)
