from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
INK = "17324D"
BLUE = "1E6A8D"
LIGHT_BLUE = "E8F1F5"
LIGHT_GRAY = "F3F5F7"
MUTED = "5E6B75"
WHITE = "FFFFFF"


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: str | None = None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, "Aptos", 9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_document(doc: Document, preset: str, running_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    body = doc.styles["Normal"]
    body.font.name = "Aptos"
    body._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    body._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(8 if preset == "narrative" else 6)
    body.paragraph_format.line_spacing = 1.28 if preset == "narrative" else 1.2
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if preset == "narrative" else WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, INK, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5 if preset == "narrative" else 0.375)
        style.paragraph_format.first_line_indent = Inches(-0.25 if preset == "narrative" else -0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(running_title)
    set_run_font(run, "Aptos", 9, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])


def add_cover(doc: Document, title: str, subtitle: str, metadata: list[str], guide: bool):
    for _ in range(4 if not guide else 2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("HOTEL SAAS PLATFORM")
    set_run_font(r, "Aptos", 11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    set_run_font(r, "Aptos Display", 25 if not guide else 23, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_run_font(r, "Aptos", 13, italic=True, color=MUTED)
    for item in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, "Aptos", 10.5, color=INK)
    doc.add_page_break()


def add_inline_runs(paragraph, text: str):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Cascadia Mono", 9.5, color=INK)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "Aptos", 11, bold=True, color=INK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Aptos", 11, color="20252A")


def parse_table(lines: list[str], start: int):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = [9360 // col_count] * col_count
    widths[-1] += 9360 - sum(widths)
    if col_count == 2:
        widths = [2700, 6660]
    elif col_count == 3:
        widths = [2100, 3000, 4260]
    elif col_count == 4:
        widths = [1500, 2200, 2860, 2800]
    for row_index, values in enumerate(rows):
        for col_index in range(col_count):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1
            text = values[col_index] if col_index < len(values) else ""
            run = paragraph.add_run(text.replace("`", ""))
            set_run_font(run, "Aptos", 9.2, bold=row_index == 0,
                         color=WHITE if row_index == 0 else "20252A")
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_pr.append(repeat_header)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, code_lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shading)
    run = p.add_run("\n".join(code_lines))
    set_run_font(run, "Cascadia Mono", 8.6, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_markdown(source: Path, output: Path, preset: str, cover: dict):
    lines = source.read_text(encoding="utf-8").splitlines()
    start_marker = cover["content_start"]
    start_index = next((index for index, line in enumerate(lines) if line.strip() == start_marker), 0)
    lines = lines[start_index:]
    doc = Document()
    configure_document(doc, preset, cover["running_title"])
    add_cover(doc, cover["title"], cover["subtitle"], cover["metadata"], cover["guide"])

    idx = 0
    in_code = False
    code_lines: list[str] = []
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and text in (cover["title"], cover["subtitle"]):
                idx += 1
                continue
            doc.add_heading(text, level=min(level, 3))
            idx += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            idx += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, numbered.group(1))
            idx += 1
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, stripped.replace("  ", " "))
        idx += 1

    doc.core_properties.title = cover["title"]
    doc.core_properties.subject = cover["subtitle"]
    doc.core_properties.author = "Hotel SaaS Platform"
    doc.core_properties.keywords = "hotel, saas, multi-tenant, rbac, booking, pms"
    doc.save(output)
    print(output)


def main():
    render_markdown(
        ROOT / "graduation-project-report.md",
        ROOT / "graduation-project-report.docx",
        "narrative",
        {
            "running_title": "Báo cáo đồ án - Hotel SaaS Platform",
            "title": "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP",
            "subtitle": "Nền tảng SaaS quản lý và đặt phòng khách sạn đa tenant",
            "metadata": [
                "Sinh viên: [HỌ VÀ TÊN SINH VIÊN] | Mã SV: [MÃ SINH VIÊN]",
                "Giảng viên hướng dẫn: [HỌ VÀ TÊN GIẢNG VIÊN]",
                "[TÊN TRƯỜNG - KHOA] | Năm học 2025-2026",
            ],
            "guide": False,
            "content_start": "## Tóm tắt",
        },
    )
    render_markdown(
        ROOT / "operations-guide.md",
        ROOT / "operations-guide.docx",
        "compact",
        {
            "running_title": "Sổ tay vận hành - Hotel SaaS Platform",
            "title": "SỔ TAY CÀI ĐẶT VÀ VẬN HÀNH",
            "subtitle": "Hotel SaaS Platform",
            "metadata": [
                "Phiên bản 1.0 | Cập nhật 21/08/2026",
                "Dành cho quản trị hệ thống, SuperAdmin và nhân sự khách sạn",
            ],
            "guide": True,
            "content_start": "# 1. Mục đích và nguyên tắc",
        },
    )


if __name__ == "__main__":
    main()
